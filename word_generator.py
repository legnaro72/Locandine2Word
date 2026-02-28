"""
Generatore di documenti Word con locandine ordinate cronologicamente
"""
import json
import os
import dateparser
import re
import streamlit as st
from typing import List, Dict
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document
from datetime import datetime


class WordGenerator:
    def __init__(self, template_path: str = None):
        """
        Inizializza il generatore
        Se template_path è None, crea un documento nuovo
        """
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
            self._setup_default_styles()
    
    def _setup_default_styles(self):
        """Configura gli stili di default del documento"""
        # Imposta margini
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def add_event_entry(self, event_data: Dict, image_path: str, mode: str = "standard", show_borders: bool = True):
        """
        Aggiunge una singola entry evento al documento
        Formato: Tabella 1x2 (immagine a sinistra, testo o immagine a destra)
        Se l'evento è scaduto, aggiunge 'completed.jpg' nella colonna di destra.
        """
        # Crea tabella 1 riga x 2 colonne
        table = self.doc.add_table(rows=1, cols=2)
        
        # Gestione bordi
        if show_borders:
            table.style = 'Table Grid'
        else:
            table.style = 'Normal Table'
        
        # Imposta larghezza colonne (40% immagine, 60% testo/immagine)
        table.columns[0].width = Inches(3.25)
        table.columns[1].width = Inches(3.25)
        
        # Cella Sinistra (Sempre Immagine)
        left_cell = table.rows[0].cells[0]
        self._insert_image(left_cell, image_path, width=Inches(2.8))

        # Cella Destra (Testo descrittivo)
        right_cell = table.rows[0].cells[1]
        self._insert_text_details(right_cell, event_data)
        
        # LOGICA SCADUTI: Se l'evento è scaduto, aggiungi l'immagine 'completed.jpg'
        now = datetime.now()
        ev_date = self.get_sort_date(event_data)
        if ev_date != datetime.max and ev_date.date() < now.date():
            completed_img = "completed.jpg"
            if os.path.exists(completed_img):
                # Aggiungi un nuovo paragrafo per l'immagine completata
                p_comp = right_cell.add_paragraph()
                p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_comp = p_comp.add_run()
                # Dimensione leggermente ridotta per stare sotto il testo
                run_comp.add_picture(completed_img, width=Inches(1.5))
        
        # LOGICA NEW: Se l'evento è contrassegnato come NEW, aggiungi l'immagine 'new.jpg'
        if event_data.get('is_new'):
            new_img_path = "new.jpg"
            if os.path.exists(new_img_path):
                p_new = right_cell.add_paragraph()
                p_new.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_new = p_new.add_run()
                run_new.add_picture(new_img_path, width=Inches(1.5))
        # Impedisci che la riga si spezzi tra due pagine
        table.rows[0].allow_break_across_pages = False

    def _insert_image(self, cell, image_path, width=Inches(2.5), max_height=Inches(3.8)):
        """Helper per inserire un'immagine in una cella, limitando sia larghezza che altezza"""
        if os.path.exists(image_path):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            
            # Calcola dimensioni reali per evitare immagini troppo alte
            try:
                from PIL import Image as PILImage
                with PILImage.open(image_path) as img:
                    img_w, img_h = img.size
                    # Rapporto aspetto: se l'immagine è più alta che larga,
                    # alla larghezza desiderata potrebbe superare max_height
                    aspect_ratio = img_h / img_w  # >1 = portrait
                    result_height_inches = (width / 914400) * aspect_ratio  # width è in EMU, 914400 EMU = 1 inch
                    max_height_inches = max_height / 914400
                    
                    if result_height_inches > max_height_inches:
                        # Troppo alta: limita all'altezza massima
                        run.add_picture(image_path, height=max_height)
                    else:
                        run.add_picture(image_path, width=width)
            except Exception:
                # Fallback: usa solo la larghezza
                run.add_picture(image_path, width=width)
            
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_text_details(self, cell, event_data):
        """Inserisce i dettagli testuali senza emoji (formato professionale)"""
        
        p = cell.paragraphs[0]

        # Titolo
        title_run = p.add_run(event_data.get('title', 'Evento') + '\n')
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Funzione helper per campo etichetta + valore
        def add_field(label, value):
            if value:
                label_run = p.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(11)

                value_run = p.add_run(f"{value}\n")
                value_run.font.size = Pt(11)

        # Data formattata
        if event_data.get('date'):
            date_str = event_data['date']
            try:
                # Supportiamo sia IT che EN per recuperare date in inglese
                dt = dateparser.parse(date_str, languages=['it', 'en'])
                if dt:
                    IT_MONTHS = {
                        1: "Gennaio", 2: "Febbraio", 3: "Marzo", 4: "Aprile",
                        5: "Maggio", 6: "Giugno", 7: "Luglio", 8: "Agosto",
                        9: "Settembre", 10: "Ottobre", 11: "Novembre", 12: "Dicembre"
                    }
                    date_formatted = f"{dt.day} {IT_MONTHS[dt.month]} {dt.year}"
                else:
                    date_formatted = date_str
            except:
                date_formatted = date_str
            add_field("DATA", date_formatted)

        add_field("ORARIO", event_data.get('time'))
        add_field("LUOGO", event_data.get('location'))
        add_field("PRESSO", event_data.get('venue'))
        add_field("INDIRIZZO", event_data.get('address'))

        # Descrizione rimosssa su richiesta utente
        pass

    
    @staticmethod
    def get_sort_date(event: Dict) -> datetime:
        """Helper statico per ottenere la data datetime da un evento con cache Streamlit."""
        d_str = event.get('date', '').strip()
        if not d_str:
            return datetime.max
            
        return WordGenerator._cached_date_parsing(d_str)

    @staticmethod
    @st.cache_data(show_spinner=False)
    def _cached_date_parsing(d_str: str) -> datetime:
        """Parsing effettivo della data con cache per massime prestazioni."""
        try:
            dt = dateparser.parse(d_str, languages=['it', 'en'])
            if dt:
                return dt
        except:
            pass
        return datetime.max # Fallback in fondo

    
    _city_fallback_cache = None

    @staticmethod
    def _get_city_fallback_dict():
        """Carica il dizionario fallback da JSON se esiste, altrimenti usa quello vuoto."""
        if WordGenerator._city_fallback_cache is not None:
            return WordGenerator._city_fallback_cache
        
        json_path = "CITY_FALLBACK.json"
        if os.path.exists(json_path):
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    WordGenerator._city_fallback_cache = json.load(f)
            except:
                WordGenerator._city_fallback_cache = {}
        else:
            WordGenerator._city_fallback_cache = {}
        
        return WordGenerator._city_fallback_cache

    @classmethod
    def reset_city_cache(cls):
        """Forza il ricaricamento della cache delle città (da chiamare dopo modifiche al JSON)"""
        cls._city_fallback_cache = None

    @staticmethod
    def get_province(event: Dict) -> str:
        """Determina la provincia di un evento in base a indirizzo e location."""
        PROV_TO_REG = {
            'GENOVA': 'LIGURIA', 'GE': 'LIGURIA',
            'LA SPEZIA': 'LIGURIA', 'SP': 'LIGURIA',
            'SAVONA': 'LIGURIA', 'SV': 'LIGURIA',
            'IMPERIA': 'LIGURIA', 'IM': 'LIGURIA',
            'ASTI': 'PIEMONTE', 'AT': 'PIEMONTE', 'ALESSANDRIA': 'PIEMONTE', 'AL': 'PIEMONTE',
            'MASSA': 'TOSCANA', 'MS': 'TOSCANA', 'MASSA CARRARA': 'TOSCANA', 'CARRARA': 'TOSCANA', 
            'LUCCA': 'TOSCANA', 'LU': 'TOSCANA'
        }
        PROV_NORM = {
            'GE': 'GENOVA', 'SP': 'LA SPEZIA', 'SV': 'SAVONA', 'IM': 'IMPERIA', 
            'AL': 'ALESSANDRIA', 'AT': 'ASTI',
            'LU': 'LUCCA',
            'MS': 'MASSA', 'MASSA CARRARA': 'MASSA', 'CARRARA': 'MASSA'
        }
        
        # Carica il dizionario dinamico dal file JSON
        CITY_FALLBACK = WordGenerator._get_city_fallback_dict()

        # Normalizzazione avanzata: rimuove spazi multipli
        raw_loc = event.get('location', '').strip().upper()
        loc = " ".join(raw_loc.split())
        
        addr = event.get('address', '').strip().upper()
        
        prov_found = None
        
        # 1. Cerca Province note nell'indirizzo (check suffisso robusto)
        sorted_provs = sorted(PROV_TO_REG.keys(), key=len, reverse=True)
        
        if addr:
            # Pulisce l'indirizzo da caratteri non alfanumerici finali
            addr_clean = re.sub(r'[^A-Z0-9]+$', '', addr)
            for p in sorted_provs:
                if addr_clean.endswith(p):
                    # Verifica che prima ci sia un separatore o spazio
                    suffix_start = len(addr_clean) - len(p)
                    if suffix_start == 0 or not addr_clean[suffix_start-1].isalnum():
                        prov_found = PROV_NORM.get(p, p)
                        break
        
        # 2. Cerca nel dizionario fallback (con location normalizzata)
        if not prov_found:
            prov_found = CITY_FALLBACK.get(loc)

        # 3. Cerca se la location stessa è una provincia
        if not prov_found:
             if loc in PROV_TO_REG:
                 prov_found = PROV_NORM.get(loc, loc)
        
        return prov_found or "ALTRO"

    def generate_from_data(self, events: List[Dict], output_path: str, mode: str = "standard", show_borders: bool = False, skip_stats: bool = False):
        """
        Genera il documento Word completo:
        1. Pagina Statistiche & Titolo (se skip_stats=False)
        2. Eventi (Standard o Minimal)
        3. Firma (se esiste)
        """
        # Creiamo un nuovo documento pulito
        self.doc = Document()
        self._setup_default_styles()

        # Aggiungi ogni evento
        sorted_events = sorted(events, key=self.get_sort_date)

        # 1. TITOLO E STATISTICHE
        if not skip_stats:
            title_text = "Eventi e Locandine"
            title = self.doc.add_heading(title_text, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = self.doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            now_str = datetime.now().strftime('%d/%m/%Y %H:%M')
            run = date_para.add_run(f"Documento generato il: {now_str}")
            run.font.size = Pt(10)
            run.font.italic = True
            
            self.doc.add_paragraph()

            st_h = self.doc.add_paragraph()
            st_h.add_run("Riepilogo Dati:").bold = True
            
            total_ev = len(sorted_events)
            self.doc.add_paragraph(f"Totale Locandine caricate: {total_ev}", style='List Bullet')
            
            PROV_TO_REG = {
                'GENOVA': 'LIGURIA', 'LA SPEZIA': 'LIGURIA', 'SAVONA': 'LIGURIA', 'IMPERIA': 'LIGURIA',
                'MASSA': 'TOSCANA', 'LUCCA': 'TOSCANA',
                'ASTI': 'PIEMONTE', 'ALESSANDRIA': 'PIEMONTE'
            }

            stats = {
                'LIGURIA': {'total': 0, 'provinces': {'GENOVA': 0, 'LA SPEZIA': 0, 'SAVONA': 0, 'IMPERIA': 0}},
                'TOSCANA': {'total': 0, 'provinces': {'MASSA': 0, 'LUCCA': 0}},
                'PIEMONTE': {'total': 0, 'provinces': {'ALESSANDRIA': 0, 'ASTI': 0}},
                'ALTRO': {'total': 0, 'cities': {}}
            }

            for ev in sorted_events:
                prov_found = self.get_province(ev)
                loc = ev.get('location', 'N/D').strip().upper()
                
                reg = PROV_TO_REG.get(prov_found)
                if reg:
                    stats[reg]['total'] += 1
                    if prov_found in stats[reg]['provinces']:
                        stats[reg]['provinces'][prov_found] += 1
                else:
                    stats['ALTRO']['total'] += 1
                    city_key = loc if loc else 'N/D'
                    stats['ALTRO']['cities'][city_key] = stats['ALTRO']['cities'].get(city_key, 0) + 1

            reg_parts = []
            for r in ['LIGURIA', 'TOSCANA', 'PIEMONTE']:
                if stats[r]['total'] > 0:
                    reg_parts.append(f"{r} ({stats[r]['total']})")
            if stats['ALTRO']['total'] > 0:
                reg_parts.append(f"ALTRO ({stats['ALTRO']['total']})")
            
            self.doc.add_paragraph(f"Distribuzione per Regione: {', '.join(reg_parts)}", style='List Bullet')

            prov_parts = []
            for r in ['LIGURIA', 'TOSCANA', 'PIEMONTE']:
                for p, count in stats[r]['provinces'].items():
                    if count > 0:
                        prov_parts.append(f"{p} ({count})")
            
            if stats['ALTRO']['total'] > 0:
                altro_cities = ", ".join([f"{c} ({n})" for c, n in sorted(stats['ALTRO']['cities'].items())])
                prov_parts.append(f"ALTRO [{altro_cities}]")
            
            if prov_parts:
                self.doc.add_paragraph(f"Distribuzione per Provincia: {', '.join(prov_parts)}", style='List Bullet')

            locations = {}
            for ev in sorted_events:
                loc = ev.get('location', 'N/D').strip().upper()
                locations[loc] = locations.get(loc, 0) + 1
            
            loc_str = ", ".join([f"{loc} ({count})" for loc, count in sorted(locations.items())])
            self.doc.add_paragraph(f"Dettaglio per Località: {loc_str}", style='List Bullet')

        # 2. ELENCO EVENTI
        # Usiamo page_break_before sui separatori invece di add_page_break()
        # per evitare paragrafi-fantasma che creano pagine bianche
        if mode == "standard":
            # Modalità Standard: 2 eventi per pagina, processati a coppie
            for page_idx in range(0, len(sorted_events), 2):
                # Separatore pagina: prima di ogni coppia
                # (dopo stats, oppure tra le coppie di eventi)
                if page_idx > 0 or not skip_stats:
                    sep = self.doc.add_paragraph()
                    sep.paragraph_format.page_break_before = True
                    sep.paragraph_format.space_before = Pt(0)
                    sep.paragraph_format.space_after = Pt(0)
                
                # Evento 1
                ev1 = sorted_events[page_idx]
                self.add_event_entry(ev1, ev1.get('image_path', ''), mode=mode, show_borders=show_borders)
                
                # Evento 2 (se esiste)
                if page_idx + 1 < len(sorted_events):
                    spacer = self.doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(6)
                    spacer.paragraph_format.space_after = Pt(6)
                    ev2 = sorted_events[page_idx + 1]
                    self.add_event_entry(ev2, ev2.get('image_path', ''), mode=mode, show_borders=show_borders)
        else:
            # Modalità Minimal: griglia 2x2 (4 eventi per pagina)
            page_size = 4
            for page_start in range(0, len(sorted_events), page_size):
                if page_start > 0 or not skip_stats:
                    sep = self.doc.add_paragraph()
                    sep.paragraph_format.page_break_before = True
                    sep.paragraph_format.space_before = Pt(0)
                    sep.paragraph_format.space_after = Pt(0)
                
                page_events = sorted_events[page_start:page_start + page_size]
                self.add_minimal_grid_page(page_events, show_borders=show_borders)
        
        # 3. FIRMA
        firma_path = "firmaComitato.docx"
        logo_path = "LogoNOConfiniTrasparente.png"
        
        if os.path.exists(firma_path):
            sep = self.doc.add_paragraph()
            sep.paragraph_format.page_break_before = True
            sep.paragraph_format.space_before = Pt(0)
            sep.paragraph_format.space_after = Pt(0)
            self._append_external_doc(firma_path)
            
            if os.path.exists(logo_path):
                p_logo = self.doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_logo = p_logo.add_run()
                run_logo.add_picture(logo_path, width=Inches(2.0))

        # Salva documento
        self.doc.save(output_path)
        return output_path

    def _append_external_doc(self, file_path):
        """Tenta di appendere il contenuto di un altro file docx"""
        try:
            external_doc = Document(file_path)
            for element in external_doc.element.body:
                self.doc.element.body.append(element)
        except Exception as e:
            # Fallback se l'append diretto fallisce (es. file corrotto)
            p = self.doc.add_paragraph()
            p.add_run(f"[Errore caricamento documento esterno {file_path}: {e}]").italic = True

    def add_minimal_grid_page(self, page_events, show_borders=False):
        """
        Aggiunge una pagina in modalità minimal con griglia 2x2:
        - Riga 1: evento 1 (sx) + evento 2 (dx)
        - Riga 2: evento 3 (sx) + evento 4 (dx)
        Ogni cella contiene Titolo + Immagine.
        page_events: lista da 1 a 4 eventi per questa pagina.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Calcola numero righe necessarie (1 o 2)
        num_rows = 2 if len(page_events) > 2 else 1
        table = self.doc.add_table(rows=num_rows, cols=2)
        
        if show_borders:
            table.style = 'Table Grid'
        else:
            table.style = 'Normal Table'
        
        # Larghezze colonne: dividono equamente la pagina
        # (pagina ~7" di testo utile con margini 0.75")
        table.columns[0].width = Inches(3.25)
        table.columns[1].width = Inches(3.25)
        
        # Impedisci che le righe si spezzino tra due pagine
        for row in table.rows:
            row.allow_break_across_pages = False
            # Altezza righe uniforme: metà pagina utile (~4.5" con margini 0.5")
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(4.5 * 1440)))  # 4.5 pollici in twip
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr.append(trHeight)
        
        # Riempi le celle con gli eventi
        # Layout: [0]=riga0-col0, [1]=riga0-col1, [2]=riga1-col0, [3]=riga1-col1
        cell_positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
        
        for idx, event in enumerate(page_events):
            if idx >= len(cell_positions):
                break
            row_idx, col_idx = cell_positions[idx]
            cell = table.rows[row_idx].cells[col_idx]
            # Allineamento verticale centrato per distribuzione uniforme
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._insert_minimal_content(cell, event)

    def add_minimal_grid_row(self, event1, event2, show_borders=False):
        """
        Aggiunge una riga in modalità minimal (backward compatibility):
        Colonna 1: Titolo + Immagine evento 1
        Colonna 2: Titolo + Immagine evento 2 (se presente)
        """
        events = [event1]
        if event2:
            events.append(event2)
        self.add_minimal_grid_page(events, show_borders=show_borders)

    def _insert_minimal_content(self, cell, event_data):
        """Inserisce Titolo (con Ora) + Immagine in una cella (modalità minimal)"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Riduce il padding interno della cella per massimizzare lo spazio
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), '30')  # ~0.02" di margine
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Spaziatura minima del paragrafo titolo
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        
        # Titolo (Grassetto, dimensione contenuta)
        # Formato: DATA - ORA - LUOGO
        title = event_data.get('title', 'Evento').upper()
        time = event_data.get('time', '').strip()
        
        if time and time not in title:
            # Ricostruzione titolo con ora se non presente
            if " - " in title:
                parts = title.split(" - ", 1)
                title_text = f"{parts[0]} - {time} - {parts[1]}"
            else:
                title_text = f"{title} - {time}"
        else:
            title_text = title

        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(8)
        # Forza il titolo a stare insieme al paragrafo successivo (l'immagine)
        p.paragraph_format.keep_with_next = True
        
        # Immagine (Sotto il titolo)
        img_path = event_data.get('image_path', '')
        
        # LOGICA SCADUTI MINIMAL: Se scaduto, mostra completed.jpg invece della locandina
        now = datetime.now()
        ev_date = WordGenerator.get_sort_date(event_data)
        if ev_date != datetime.max and ev_date.date() < now.date():
            if os.path.exists("completed.jpg"):
                img_path = "completed.jpg"

        if img_path and os.path.exists(img_path):
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(0)
            run_img = p_img.add_run()
            # Larghezza calcolata per stare bene in 2x2
            run_img.add_picture(img_path, width=Inches(2.6))
    
    def load_events_from_json(self, json_path: str) -> List[Dict]:
        """Carica eventi dal file JSON"""
        if os.path.exists(json_path):
            with open(json_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def save_events_to_json(self, events: List[Dict], json_path: str):
        """Salva eventi nel file JSON"""
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(events, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    # Test
    generator = WordGenerator()
    print("Word Generator inizializzato correttamente!")